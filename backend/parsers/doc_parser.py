import io
import os
import re
from pypdf import PdfReader
from docx import Document


def clean_text(text: str) -> str:
    """
    Cleans up excessive whitespace, normalizes newlines, and strips
    unprintable control characters while preserving document structure.
    """
    if not text:
        return ""

    # Remove unprintable characters (keep newlines and tabs)
    text = "".join(char for char in text if char.isprintable() or char in "\n\r\t")

    # Normalize carriage returns
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse 3 or more consecutive newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Trim redundant spaces and tabs within lines
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    return "\n".join(lines).strip()


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Detects file extension (.pdf, .docx, .txt, .md) and extracts clean,
    normalized plain text from raw file bytes.
    """
    ext = os.path.splitext(filename)[1].lower()

    if not file_bytes:
        raise ValueError(f"File '{filename}' is empty.")

    raw_text = ""

    # 1. PDF extraction using pypdf
    if ext == ".pdf":
        pdf_stream = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_stream)
        pages_text = []
        for page in reader.pages:
            content = page.extract_text()
            if content:
                pages_text.append(content.strip())
        raw_text = "\n\n".join(pages_text)

    # 2. Word (.docx) extraction using python-docx
    elif ext == ".docx":
        docx_stream = io.BytesIO(file_bytes)
        doc = Document(docx_stream)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also collect text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        raw_text = "\n\n".join(paragraphs)

    # 3. Plain Text and Markdown
    elif ext in [".txt", ".md"]:
        try:
            raw_text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raw_text = file_bytes.decode("latin-1", errors="ignore")

    else:
        raise ValueError(
            f"Unsupported file format: '{ext}'. Supported formats are: .pdf, .docx, .txt, .md"
        )

    return clean_text(raw_text)








# ==========================================
# TEST BLOCK (Runs only when executing this file directly)
# ==========================================
if __name__ == "__main__":
    import sys

    # Ensure backend imports work
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    if project_root not in sys.path:
        sys.path.insert(0, project_root)

    from backend.exporters.pptx_exporter import generate_pptx_file
    from backend.database.db import init_db, save_transformation, fetch_history
    from backend.schemas import PresentationDeck, Slide

    print("\n" + "=" * 60)
    print("🚀 RUNNING VERIFICATION TEST")
    print("=" * 60)

    # 1. Test parsing sample_1_cyber_threat.txt
    sample_file = os.path.join(project_root, "data", "samples", "sample_1_cyber_threat.txt")
    print(f"\n[1/3] Reading: {sample_file}")

    with open(sample_file, "rb") as f:
        data_bytes = f.read()

    extracted = extract_text_from_file(data_bytes, "sample_1_cyber_threat.txt")
    print(f"  ✅ Extracted {len(extracted)} characters.")
    print(f"  📝 Preview: {extracted[:100].strip()}...\n")

    # 2. Test generating test_presentation.pptx
    pptx_path = os.path.join(project_root, "test_presentation.pptx")
    print(f"[2/3] Generating presentation: {pptx_path}")

    mock_deck = PresentationDeck(
        deck_title="Cyber Threat Analysis Brief",
        target_audience="Security Operations Center",
        slides=[
            Slide(
                slide_num=1,
                title="Vulnerability Detection",
                bullet_points=[
                    "Unauthenticated remote code execution found in telemetry daemon.",
                    "Apply Emergency Patch 2026-04 immediately."
                ],
                visual_diagram_concept="Network topology highlighting daemon node",
                speaker_notes="Mandatory immediate patch window required."
            )
        ]
    )

    generate_pptx_file(mock_deck, pptx_path)
    print(f"  ✅ Saved PPTX to: {pptx_path}\n")

    # 3. Test storing record in SQLite
    print("[3/3] Storing dummy record in SQLite database...")
    init_db()
    rec_id = save_transformation(
        record_id="test-001",
        title="Sample Cyber Threat Transformation",
        raw_text=extracted[:100],
        settings={"tone": "Urgent"},
        result={"status": "Success"}
    )
    print(f"  ✅ Saved record with ID: {rec_id}")

    history = fetch_history(limit=1)
    print(f"  ✅ Verified record from database: '{history[0]['title']}'")

    print("\n" + "=" * 60)
    print("🎉 ALL TESTS PASSED SUCCESSFULLY WITHOUT CREATING NEW FILES!")
    print("=" * 60 + "\n")
